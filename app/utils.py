from __future__ import annotations
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document
import os

def extract_text_from_pdf(path: str) -> str:
    try:
        return pdf_extract_text(path) or ""
    except Exception:
        return ""

def extract_text_from_docx(path: str) -> str:
    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception:
        return ""

def extract_text_auto(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    if ext in (".docx", ):
        return extract_text_from_docx(path)
    return ""
